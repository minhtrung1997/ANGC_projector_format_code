#!/usr/bin/env python
# coding: utf-8

import docx
import argparse

def main():
    Parser = argparse.ArgumentParser(description='Convert the heading lines of a docx file to heading 2 format')
    Parser.add_argument('-d','--docx_path', metavar='docx_path', type=str, help='The docx file', required=True)
    Parser.add_argument('-o','--output_path', metavar='output_path', type=str, help='The output docx file, if not specified, the input file will be overwritten', required=False)
    args = Parser.parse_args()

    docx_path = args.docx_path
    if args.output_path:
        output_path = args.output_path
    else:
        output_path = docx_path
    # Parse the file in docx format
    doc = docx.Document(docx_path)

    # Get the text of the file
    text = []
    for para in doc.paragraphs:
        text.append(para.text)

    # Find the heading lines of the text which are in bold and contain the string '(O)' or '(W)' or '(P)' then convert them to heading 2 format
    for i in range(len(text)):
        if text[i].endswith('(O)') or text[i].endswith('(W)') or text[i].endswith('(P)'):
            # remove the break line character at the beginning of the line (if any) but keep the break line
            if text[i].startswith('\n'):
                doc.paragraphs[i].text = text[i][1:]
            else:
                doc.paragraphs[i].text = text[i]
            doc.paragraphs[i].style = 'Heading 2'
            if i > 1:
                doc.paragraphs[i-1].text = doc.paragraphs[i-1].text + ' (Finished the song)'
                # Add page break at the end of the previous line
                doc.paragraphs[i-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
            # print out the heading lines
            print(doc.paragraphs[i].text)

    # Find the line contain 'INDEX' case sensitive string and remove it and all the lines after it
    for i in range(len(text)):
        if text[i].startswith('INDEX'):
            for j in range(i, len(text)):
                print(f'Remove line "{doc.paragraphs[j].text}"')
                doc.paragraphs[j].clear()
            break

    # Save the file
    doc.save(output_path)

if __name__ == '__main__':
    main()
