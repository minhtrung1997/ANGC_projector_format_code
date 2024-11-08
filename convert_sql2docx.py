# This script will parse an SQL file and convert it into a pandas dataframe
# It then reformat the dataframe and save it as a docx files

import pandas as pd
import argparse
import docx
import sqlite3
from local_function import convert_contents

conn = sqlite3.connect('download_db/Dedupe_Pooled_Database_2023_curated.db')

sql_query = pd.read_sql_query('''
                               SELECT
                               title1,
                               title2,
                               contents,
                               sequence
                               FROM items
                               ''', conn)

df = pd.DataFrame(sql_query, columns = ['title1', 'title2', 'contents'])
# Remove row if the contents column is less than 15 characters
df = df[df['contents'].str.len() > 15]
# reset the index
df = df.reset_index(drop=True)
# We will use the convert_contents.main function to parse the contents column
# The function will split the contents into content1 and content2 and return them as a string
# Create the new columns
df['content1'] = ''
df['content2'] = ''
for i in range(len(df)):
    # check if contents has [region 2] marker if not then content1 equals contents and content2 equals ''
    if '[region 2]' not in df.loc[i, 'contents']:
        # df['content1'][i] = df['contents'][i]
        # df['content2'][i] = ''
        df.loc[i, 'content1'] = df.loc[i, 'contents']
        df.loc[i, 'content2'] = ''
    else:
        # df['content1'][i], df['content2'][i] = convert_contents.main(df['contents'][i])
        df.loc[i, 'content1'], df.loc[i, 'content2'] = convert_contents.main(df.loc[i, 'contents'])

# drop the contents column
df = df.drop(columns = ['contents'])
# Sort df by title1 alphabetically
df = df.sort_values(by=['title1'])
df = df.reset_index(drop=True)

# Now we need to reformat the dataframe and write it to a docx file
# The title1 will be the heading
# below is the title2 (not a heading) and content1 and content2
# Contents will be placed in a double column table format
# Content1 will be in the left column and content2 will be in the right column

doc = docx.Document()

from docx.shared import RGBColor

def add_text_with_color(cell, text):
    parts = text.split('[')
    for part in parts:
        if ']' in part:
            before_bracket, after_bracket = part.split(']', 1)
            run = cell.paragraphs[0].add_run('[' + before_bracket + ']')
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            run.font.bold = True
            run = cell.paragraphs[0].add_run(after_bracket)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Reset to black color
            run.font.bold = False
        else:
            run = cell.paragraphs[0].add_run(part)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black color for non-bracket text
            run.font.bold = False

for i in range(len(df)):
    doc.add_heading(f"Song {i+1}: {df['title1'][i]}", level=1)
    doc.add_paragraph(df['title2'][i])
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = df['content1'][i]
    # hdr_cells[1].text = df['content2'][i]
    add_text_with_color(hdr_cells[0], df['content1'][i])
    add_text_with_color(hdr_cells[1], df['content2'][i])

    doc.add_paragraph()

# make sure the output folder exists
import os
if not os.path.exists('output'):
    os.makedirs('output')
doc.save('output/Danh sách bài hát ANGC.docx')

from docx import Document
from docx.oxml.shared import OxmlElement, qn
import lxml.etree
from docx.oxml import parse_xml
import subprocess

def add_table_of_contents(document):
    """Inserts a basic table of contents field into the given Word document."""
    doc = document

    # Add "Mục Lục" heading
    heading = doc.add_heading('Mục Lục', level=1)._element
    doc.element.body.insert(0, heading)

    # Insert TOC field with instructions (Modified section)
    toc = doc.add_paragraph()
    run = toc.add_run()

    # Add namespace declarations to each element where 'w' is used
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    fldChar.set(qn('w:dirty'), 'true')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = r'TOC \o "1-3" \h \z \u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    # Insert TOC at the first position
    toc = toc._element
    doc.element.body.insert(1, toc)

    # Add a page break after the TOC
    page_break = doc.add_page_break()._element
    doc.element.body.insert(2, page_break)

# Create or open an existing document
doc = Document('output/Danh sách bài hát ANGC.docx')
add_table_of_contents(doc)
doc.save('output/Danh sách bài hát ANGC.docx')
print("Finished creating the document")
